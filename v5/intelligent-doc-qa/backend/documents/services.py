"""
Service for processing documents: text extraction, chunking, and embedding.
"""

import os
import io
from typing import List, Tuple, Optional
from datetime import datetime

import PyPDF2
import pdfplumber
from docx import Document as DocxDocument
from PIL import Image
import google.generativeai as genai
from google.cloud import vision
from django.conf import settings
from langchain.text_splitter import RecursiveCharacterTextSplitter

from documents.models import Document, DocumentChunk


class DocumentProcessor:
    """Handles document text extraction, chunking, and embedding generation."""
    
    def __init__(self):
        """Initialize the document processor with API keys and settings."""
        genai.configure(api_key=settings.GOOGLE_API_KEY)
        self.chunk_size = settings.CHUNK_SIZE
        self.chunk_overlap = settings.CHUNK_OVERLAP
        self.embedding_model = settings.GEMINI_EMBEDDING_MODEL
        
        # Initialize text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
    
    def process_document(self, document: Document) -> bool:
        """
        Process a document: extract text, chunk it, generate embeddings.
        
        Args:
            document: Document model instance
            
        Returns:
            bool: True if successful, False otherwise
        """
        try:
            # Update status to PROCESSING
            document.status = Document.Status.PROCESSING
            document.save()
            
            # Extract text based on file type
            text, num_pages = self._extract_text(document)
            
            if not text or not text.strip():
                document.status = Document.Status.FAILED
                document.error_message = "No text extracted from document"
                document.save()
                return False
            
            # Split into chunks
            chunks = self._chunk_text(text)
            
            # Generate embeddings and save chunks
            self._save_chunks(document, chunks)
            
            # Update document status
            document.status = Document.Status.READY
            document.num_chunks = len(chunks)
            document.num_pages = num_pages
            document.processed_at = datetime.now()
            document.error_message = None
            document.save()
            
            return True
            
        except Exception as e:
            document.status = Document.Status.FAILED
            document.error_message = str(e)
            document.save()
            return False
    
    def _extract_text(self, document: Document) -> Tuple[str, int]:
        """
        Extract text from document based on file type.
        
        Args:
            document: Document model instance
            
        Returns:
            Tuple of (extracted_text, num_pages)
        """
        file_path = document.file_path.path
        file_type = document.file_type.lower()
        
        if file_type == 'pdf':
            return self._extract_from_pdf(file_path)
        elif file_type == 'docx':
            return self._extract_from_docx(file_path)
        elif file_type == 'txt':
            return self._extract_from_txt(file_path)
        elif file_type in ['jpg', 'jpeg', 'png']:
            return self._extract_from_image(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    def _extract_from_pdf(self, file_path: str) -> Tuple[str, int]:
        """Extract text from PDF file."""
        text_parts = []
        num_pages = 0
        
        try:
            # Try pdfplumber first (better for complex PDFs)
            with pdfplumber.open(file_path) as pdf:
                num_pages = len(pdf.pages)
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
        except Exception as e:
            # Fallback to PyPDF2
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    num_pages = len(pdf_reader.pages)
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
            except Exception as fallback_error:
                raise Exception(f"Failed to extract PDF text: {fallback_error}")
        
        return "\n\n".join(text_parts), num_pages
    
    def _extract_from_docx(self, file_path: str) -> Tuple[str, int]:
        """Extract text from DOCX file."""
        doc = DocxDocument(file_path)
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                text_parts.append(" | ".join(row_text))
        
        num_pages = len(doc.sections)  # Approximate
        return "\n\n".join(text_parts), num_pages
    
    def _extract_from_txt(self, file_path: str) -> Tuple[str, int]:
        """Extract text from TXT file."""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            text = file.read()
        
        # Estimate pages (roughly 3000 chars per page)
        num_pages = max(1, len(text) // 3000)
        return text, num_pages
    
    def _extract_from_image(self, file_path: str) -> Tuple[str, int]:
        """Extract text from image using Google Vision API or Gemini."""
        try:
            # Try using Gemini's multimodal capabilities
            model = genai.GenerativeModel('gemini-pro-vision')
            image = Image.open(file_path)
            
            response = model.generate_content([
                "Extract all text from this image. Return only the text, nothing else.",
                image
            ])
            
            text = response.text
            return text, 1
            
        except Exception as e:
            # Fallback: try Google Cloud Vision API if available
            try:
                client = vision.ImageAnnotatorClient()
                
                with open(file_path, 'rb') as image_file:
                    content = image_file.read()
                
                image = vision.Image(content=content)
                response = client.text_detection(image=image)
                texts = response.text_annotations
                
                if texts:
                    return texts[0].description, 1
                else:
                    raise Exception("No text found in image")
                    
            except Exception as vision_error:
                raise Exception(f"Failed to extract text from image: {vision_error}")
    
    def _chunk_text(self, text: str) -> List[str]:
        """
        Split text into chunks using LangChain text splitter.
        
        Args:
            text: Full text to chunk
            
        Returns:
            List of text chunks
        """
        chunks = self.text_splitter.split_text(text)
        return chunks
    
    def _generate_embedding(self, text: str) -> List[float]:
        """
        Generate embedding for text using Gemini.
        
        Args:
            text: Text to embed
            
        Returns:
            List of floats representing the embedding
        """
        result = genai.embed_content(
            model=self.embedding_model,
            content=text,
            task_type="retrieval_document"
        )
        return result['embedding']
    
    def _save_chunks(self, document: Document, chunks: List[str]) -> None:
        """
        Save chunks with embeddings to database.
        
        Args:
            document: Document model instance
            chunks: List of text chunks
        """
        # Delete existing chunks if any
        DocumentChunk.objects.filter(document=document).delete()
        
        chunk_objects = []
        for idx, chunk_text in enumerate(chunks):
            # Generate embedding
            embedding = self._generate_embedding(chunk_text)
            
            # Create chunk object
            chunk = DocumentChunk(
                document=document,
                index=idx,
                text=chunk_text,
                embedding=embedding,
                char_count=len(chunk_text),
                token_count=len(chunk_text.split())  # Rough approximation
            )
            chunk_objects.append(chunk)
        
        # Bulk create all chunks
        DocumentChunk.objects.bulk_create(chunk_objects)
