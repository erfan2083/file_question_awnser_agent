"""
DOCX text extractor implementation.
"""
from typing import Tuple

from docx import Document as DocxDocument

from core.domain.value_objects.file_type import FileType
from core.domain.exceptions import TextExtractionError
from core.application.ports.services.text_extractor import TextExtractor


class DOCXTextExtractor(TextExtractor):
    """DOCX text extraction implementation."""

    def can_extract(self, file_type: FileType) -> bool:
        """Check if this extractor can handle DOCX files."""
        return file_type == FileType.DOCX

    def extract(self, file_path: str) -> Tuple[str, int]:
        """Extract text from DOCX file."""
        try:
            doc = DocxDocument(file_path)
            text_parts = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    text_parts.append(" | ".join(row_text))

            num_pages = len(doc.sections)  # Approximate
            return "\n\n".join(text_parts), num_pages
        except Exception as e:
            raise TextExtractionError(f"Failed to extract DOCX text: {str(e)}")
